from datetime import datetime

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# =====================================================
# CONFIGURAR ESTILOS
# =====================================================

def configure_styles(

    document

):

    normal = document.styles["Normal"]

    normal.font.name = "Calibri"

    normal.font.size = Pt(11)

    h1 = document.styles["Heading 1"]

    h1.font.name = "Calibri"

    h1.font.bold = True

    h1.font.size = Pt(20)

    h2 = document.styles["Heading 2"]

    h2.font.name = "Calibri"

    h2.font.bold = True

    h2.font.size = Pt(15)

    h3 = document.styles["Heading 3"]

    h3.font.name = "Calibri"

    h3.font.bold = True

    h3.font.size = Pt(12)

# =====================================================
# PORTADA
# =====================================================

def add_cover(

    document,

    team

):

    title = document.add_heading()

    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = title.add_run(

        "INFORME TÉCNICO"

    )

    run.bold = True

    run.font.size = Pt(28)

    document.add_paragraph()

    subtitle = document.add_paragraph()

    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = subtitle.add_run(

        "Plataforma de Inteligencia Deportiva"

    )

    run.bold = True

    run.font.size = Pt(16)

    document.add_paragraph()

    p = document.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p.add_run(

        "Equipo analizado\n"

    ).bold = True

    p.add_run(

        team

    )

    document.add_paragraph()

    p = document.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p.add_run(

        "Fecha\n"

    ).bold = True

    p.add_run(

        datetime.now().strftime(

            "%d/%m/%Y"

        )

    )

    document.add_paragraph()

    p = document.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p.add_run(

        "Analista\n"

    ).bold = True

    p.add_run(

        "David Vázquez Diéguez"

    )

    document.add_paragraph()

    p = document.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run(

        "DOCUMENTO CONFIDENCIAL"

    )

    run.italic = True

    run.font.size = Pt(10)

    document.add_page_break()

# =====================================================
# CABECERA DEL MÓDULO
# =====================================================

def add_module(

    document,

    title

):

    table = document.add_table(

        rows=1,

        cols=1

    )

    table.style = "Table Grid"

    cell = table.cell(

        0,

        0

    )

    p = cell.paragraphs[0]

    run = p.add_run(

        title.upper()

    )

    run.bold = True

    run.font.size = Pt(18)

    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.add_paragraph()


# =====================================================
# TÍTULO DE SECCIÓN
# =====================================================

def add_section(

    document,

    title

):

    p = document.add_paragraph()

    run = p.add_run(

        title

    )

    run.bold = True

    run.font.size = Pt(14)

    document.add_paragraph()


# =====================================================
# VARIABLES
# =====================================================

def add_variables(

    document,

    variables

):

    if not variables:

        return

    if isinstance(

        variables,

        list

    ):

        text = " · ".join(

            variables

        )

    else:

        text = str(

            variables

        )

    p = document.add_paragraph()

    run = p.add_run(

        "Variables utilizadas: "

    )

    run.bold = True

    p.add_run(

        text

    )

    document.add_paragraph()

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT


# =====================================================
# BLOQUE DE COMENTARIO
# =====================================================

def add_comment(

    document,

    note_type,

    comment,

    date=None

):

    icons = {

        "Fortaleza": "🟢",

        "Debilidad": "🔴",

        "Idea táctica": "🔵",

        "Dato relevante": "🟡",

        "Observación": "⚪"

    }

    icon = icons.get(

        note_type,

        "⚪"

    )

    table = document.add_table(

        rows=1,

        cols=1

    )

    table.style = "Table Grid"

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(

        0,

        0

    )

    p = cell.paragraphs[0]

    run = p.add_run(

        f"{icon} {note_type.upper()}"

    )

    run.bold = True

    run.font.size = Pt(12)

    p = cell.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    run = p.add_run(

        comment

    )

    run.font.size = Pt(11)

    if date:

        p = cell.add_paragraph()

        run = p.add_run(

            f"Actualizado: {date}"

        )

        run.italic = True

        run.font.size = Pt(8)

    document.add_paragraph()

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


# =====================================================
# CABECERA INSTITUCIONAL
# =====================================================

def add_report_header(

    document,

    team

):

    p = document.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run(

        "PLATAFORMA DE ANÁLISIS DE DATOS"

    )

    run.bold = True

    run.font.size = Pt(10)

    p = document.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run(

        f"{team} · {datetime.now().strftime('%d/%m/%Y')}"

    )

    run.font.size = Pt(9)

    document.add_paragraph()


# =====================================================
# SEPARADOR
# =====================================================

def add_separator(

    document

):

    p = document.add_paragraph()

    run = p.add_run(

        "─" * 80

    )

    run.font.size = Pt(8)

    document.add_paragraph()


# =====================================================
# PIE DEL INFORME
# =====================================================

def add_signature(

    document

):

    document.add_paragraph()

    document.add_paragraph()

    p = document.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    run = p.add_run(

        "Departamento de Análisis"

    )

    run.bold = True

    run.font.size = Pt(10)

    p = document.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p.add_run(

        "David Vázquez Diéguez"

    )

    p = document.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    run = p.add_run(

        datetime.now().strftime(

            "%d/%m/%Y"

        )

    )

    run.italic = True

    run.font.size = Pt(9)

from pathlib import Path

from docx.shared import Inches


# =====================================================
# LOGO DEL CLUB
# =====================================================

def add_logo(

    document,

    team

):

    logo = Path(

        "assets/logos"

    ) / f"{team}.png"

    if logo.exists():

        document.add_picture(

            str(logo),

            width=Inches(1.2)

        )

        document.add_paragraph()


# =====================================================
# IMAGEN DE GRÁFICO
# =====================================================

def add_chart(

    document,

    image_path,

    width=6.3

):

    image = Path(

        image_path

    )

    if image.exists():

        document.add_picture(

            str(image),

            width=Inches(width)

        )

        document.add_paragraph()


# =====================================================
# SALTO DE PÁGINA
# =====================================================

def add_page(

    document

):

    document.add_page_break()